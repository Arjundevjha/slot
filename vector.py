from langchain_ollama import OllamaEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document
import os
import docx 

# Define the path to the .docx file
docx_file = "SL Symposium Programme 2025_28 Feb (1).docx"
embeddings = OllamaEmbeddings(model="mxbai-embed-large")

db_location = "./chrome_langchain"
add_documents = not os.path.exists(db_location)

if add_documents:
    documents = []
    ids = []
    
    # Open the .docx file
    doc = docx.Document(docx_file)

    # Extract data from tables
    for table in doc.tables:
        for row in table.rows:
            text_cells = [cell.text for cell in row.cells]
            # Assuming the table structure is consistent, adjust accordingly
            if len(text_cells) == 2:  # Example: Time, Programme table [cite: 88]
                page_content = f"Time: {text_cells[0]}, Programme: {text_cells[1]}"
                document = Document(
                    page_content=page_content,
                    metadata={"source": "table"},
                )
                documents.append(document)
                ids.append(str(len(documents) -1 ))

    # Extract data from paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text.strip():  # Ignore empty paragraphs
            document = Document(
                page_content=text,
                metadata={"source": "paragraph"},
            )
            documents.append(document)
            ids.append(str(len(documents) - 1))

vector_store = Chroma(
    collection_name="Event Information",
    persist_directory=db_location,
    embedding_function=embeddings
)

if add_documents:
    vector_store.add_documents(documents=documents, ids=ids)

retriever = vector_store.as_retriever(
    search_kwargs={"k": 3}
)